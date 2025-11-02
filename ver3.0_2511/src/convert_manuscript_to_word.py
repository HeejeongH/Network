#!/usr/bin/env python3
"""
Markdown을 Word 문서로 변환하는 스크립트
Paper2_Main_Manuscript.md -> Paper2_Main_Manuscript.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def setup_styles(doc):
    """Word 문서 스타일 설정"""
    # 제목 1 스타일
    if 'Heading 1' in doc.styles:
        h1 = doc.styles['Heading 1']
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
    
    # 제목 2 스타일
    if 'Heading 2' in doc.styles:
        h2 = doc.styles['Heading 2']
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
    
    # 제목 3 스타일
    if 'Heading 3' in doc.styles:
        h3 = doc.styles['Heading 3']
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0, 0, 0)
    
    # 본문 스타일
    if 'Normal' in doc.styles:
        normal = doc.styles['Normal']
        normal.font.size = Pt(11)
        normal.font.name = 'Times New Roman'

def parse_markdown_line(line):
    """마크다운 줄 파싱"""
    line = line.rstrip()
    
    # 빈 줄
    if not line:
        return 'empty', ''
    
    # 제목 파싱
    if line.startswith('# '):
        return 'h1', line[2:]
    elif line.startswith('## '):
        return 'h2', line[3:]
    elif line.startswith('### '):
        return 'h3', line[4:]
    elif line.startswith('#### '):
        return 'h4', line[5:]
    
    # 구분선
    if line.strip() in ['---', '***', '___']:
        return 'separator', ''
    
    # 리스트
    if re.match(r'^\d+\.\s', line):
        return 'ordered_list', re.sub(r'^\d+\.\s', '', line)
    elif line.startswith('- ') or line.startswith('* '):
        return 'unordered_list', line[2:]
    
    # 일반 텍스트
    return 'text', line

def apply_inline_formatting(paragraph, text):
    """인라인 포맷팅 적용 (볼드, 이탤릭 등)"""
    # **bold** 패턴
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # *italic* 패턴
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ipart)

def convert_markdown_to_word(md_file, output_file):
    """마크다운 파일을 Word 문서로 변환"""
    print(f"Converting {md_file} to {output_file}...")
    
    # Word 문서 생성
    doc = Document()
    setup_styles(doc)
    
    # 페이지 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 마크다운 파일 읽기
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 줄별로 처리
    for i, line in enumerate(lines):
        line_type, content = parse_markdown_line(line)
        
        if line_type == 'empty':
            # 빈 줄 추가 (연속된 빈 줄은 하나만)
            if i > 0 and parse_markdown_line(lines[i-1])[0] != 'empty':
                doc.add_paragraph()
        
        elif line_type == 'h1':
            p = doc.add_heading(content, level=1)
        
        elif line_type == 'h2':
            p = doc.add_heading(content, level=2)
        
        elif line_type == 'h3':
            p = doc.add_heading(content, level=3)
        
        elif line_type == 'h4':
            # Word는 기본적으로 heading 4 지원
            p = doc.add_paragraph(content, style='Heading 4')
        
        elif line_type == 'separator':
            # 구분선은 빈 줄로
            doc.add_paragraph()
        
        elif line_type == 'ordered_list':
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, content)
        
        elif line_type == 'unordered_list':
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, content)
        
        elif line_type == 'text':
            p = doc.add_paragraph()
            apply_inline_formatting(p, content)
    
    # 파일 저장
    doc.save(output_file)
    print(f"✓ Saved to {output_file}")
    print(f"  File size: {output_file.stat().st_size / 1024:.1f} KB")

def main():
    """메인 함수"""
    # 경로 설정
    base_dir = Path(__file__).parent.parent
    manuscript_dir = base_dir / 'result' / 'manuscript'
    
    # 입력 파일
    md_file = manuscript_dir / 'Paper2_Main_Manuscript.md'
    
    # 출력 파일
    output_file = manuscript_dir / 'Paper2_Main_Manuscript.docx'
    
    if not md_file.exists():
        print(f"Error: {md_file} not found!")
        return
    
    # 변환 실행
    convert_markdown_to_word(md_file, output_file)
    
    print("\n" + "="*60)
    print("변환 완료!")
    print("="*60)
    print(f"\n파일 위치: {output_file}")
    print("\n이제 Word 파일을 열어서 확인하시면 됩니다.")
    print("필요시 Word에서 추가 서식을 조정할 수 있습니다.")

if __name__ == '__main__':
    main()
