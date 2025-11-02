#!/usr/bin/env python3
"""
SCI급 영어 논문 작성 스크립트
기존 내용을 활용하되 더 학술적이고 간결하게 리라이팅
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_sci_paper():
    """SCI급 영어 논문 작성"""
    doc = Document()
    
    # 페이지 설정
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Universal and Group-Specific Hub Foods in Dietary Co-occurrence Networks: ')
    run.bold = True
    run.font.size = Pt(14)
    run = title.add_run('A Stratified Analysis of 22,964 Korean Adults')
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Running title
    p = doc.add_paragraph()
    run = p.add_run('Running Title: ')
    run.bold = True
    p.add_run('Dietary Network Hubs Across Demographics')
    
    doc.add_paragraph()
    
    # Authors (placeholder)
    p = doc.add_paragraph()
    run = p.add_run('Authors: ')
    run.bold = True
    p.add_run('[To be filled]')
    
    doc.add_paragraph()
    
    # Abstract
    p = doc.add_heading('Abstract', level=1)
    
    # Background
    p = doc.add_paragraph()
    run = p.add_run('Background: ')
    run.bold = True
    p.add_run('Traditional dietary pattern analysis often overlooks the complex interconnections among foods and how these patterns vary across population subgroups. Network analysis offers a novel approach to identify "hub" foods that are central to dietary patterns, but few studies have examined how these hubs differ by demographic and metabolic characteristics.')
    
    # Objective
    p = doc.add_paragraph()
    run = p.add_run('Objective: ')
    run.bold = True
    p.add_run('We applied stratified co-occurrence network analysis to identify universal and group-specific hub foods across sex, age, and metabolic syndrome (MetS) status in a large Korean population.')
    
    # Methods
    p = doc.add_paragraph()
    run = p.add_run('Methods: ')
    run.bold = True
    p.add_run('Data from 22,964 adults (19-74 years) in the Korea National Health and Nutrition Examination Survey were stratified into 11 groups by sex, age (19-39, 40-59, 60-74 years), and MetS status. Co-occurrence networks were constructed for 12 food groups based on simultaneous high-consumption patterns (score ≥3 on 3- or 4-point scales). Hub foods were identified using degree, betweenness, and closeness centrality metrics.')
    
    # Results  
    p = doc.add_paragraph()
    run = p.add_run('Results: ')
    run.bold = True
    p.add_run('All networks showed identical topology (12 nodes, 20 edges, density=0.303) but differing centrality distributions. Three foods emerged as universal hubs across all 11 groups: protein foods (degree centrality 0.636-1.000), vegetables (0.455-1.000), and grain products (0.364-0.545). Group-specific hub patterns included higher centrality for sugar-sweetened beverages in young adults (0.273-0.364 vs. 0.091-0.273 in older adults) and sweet foods in females. Hub composition shifted with age, particularly in males, from sugar-sweetened beverages in youth to grain products in older adulthood.')
    
    # Conclusions
    p = doc.add_paragraph()
    run = p.add_run('Conclusions: ')
    run.bold = True
    p.add_run('Despite uniform network structures, hub food centrality varies substantially across demographic and metabolic subgroups. The protein-vegetable-grain triad represents a universal dietary core suitable for population-wide interventions, while age- and sex-specific hub patterns support tailored nutritional counseling. These findings demonstrate the value of stratified network analysis for personalized nutrition strategies.')
    
    # Keywords
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.bold = True
    p.add_run('dietary networks; co-occurrence analysis; hub foods; metabolic syndrome; stratified analysis; personalized nutrition')
    
    doc.add_page_break()
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    
    p = doc.add_paragraph(
        'Metabolic syndrome (MetS) affects one-quarter of the global adult population and substantially '
        'increases risk for cardiovascular disease and type 2 diabetes (1, 2). Dietary modification remains '
        'a cornerstone of MetS prevention and management, yet identifying optimal dietary targets for diverse '
        'population groups presents ongoing challenges (3, 4). Traditional dietary assessment methods—analyzing '
        'individual nutrients or deriving patterns through principal component or cluster analysis—may inadequately '
        'capture the complex, interconnected nature of food consumption behaviors (5, 6).'
    )
    
    p = doc.add_paragraph(
        'Network science offers a complementary framework for dietary pattern analysis by explicitly modeling '
        'foods as nodes and their co-consumption relationships as edges (7, 8). This approach reveals "hub" '
        'foods that occupy central positions within dietary networks and may serve as leverage points for behavioral '
        'interventions (9, 10). However, most dietary network studies have analyzed populations as single entities, '
        'potentially obscuring important heterogeneity in dietary patterns across demographic and clinical subgroups '
        '(11, 12). Given established differences in dietary preferences by sex and age (13, 14), and the complex '
        'bidirectional relationship between diet and metabolic health (15, 16), stratified network analysis is needed '
        'to inform personalized nutrition strategies.'
    )
    
    p = doc.add_paragraph(
        'We conducted a stratified co-occurrence network analysis of dietary patterns in 22,964 Korean adults, '
        'examining 11 groups defined by sex, age, and MetS status. Our objectives were to: (1) identify hub foods '
        'within each demographic-metabolic subgroup using multiple centrality metrics; (2) distinguish universal hubs '
        'appearing across all groups from group-specific hubs; and (3) characterize hub food transitions across the '
        'lifespan. These findings provide evidence-based targets for both population-wide and personalized dietary '
        'interventions.'
    )
    
    doc.add_page_break()
    
    # Methods
    doc.add_heading('Methods', level=1)
    
    doc.add_heading('Study Population', level=2)
    p = doc.add_paragraph(
        'We analyzed data from the Korea National Health and Nutrition Examination Survey (KNHANES), a nationally '
        'representative cross-sectional survey employing complex, stratified, multistage probability sampling. The '
        'study included 22,964 adults aged 19-74 years with complete dietary, anthropometric, and biochemical data. '
        'Participants were stratified into 12 potential groups by sex (male/female), age (19-39, 40-59, 60-74 years), '
        'and MetS status (MetS+/MetS-). The female young adult MetS+ group was excluded due to insufficient sample '
        'size (n<100), yielding 11 groups for analysis (range: 516-5,629 participants per group). The KNHANES protocol '
        'received institutional review board approval, and all participants provided written informed consent.'
    )
    
    doc.add_heading('Metabolic Syndrome Definition', level=2)
    p = doc.add_paragraph(
        'MetS was defined using modified NCEP ATP III criteria with Asian-specific waist circumference thresholds (17). '
        'Participants meeting ≥3 of five criteria were classified as MetS+: (1) waist circumference ≥90 cm (men) or '
        '≥85 cm (women); (2) triglycerides ≥150 mg/dL or medication use; (3) HDL-cholesterol <40 mg/dL (men) or '
        '<50 mg/dL (women); (4) blood pressure ≥130/85 mmHg or medication use; (5) fasting glucose ≥100 mg/dL or '
        'medication use.'
    )
    
    doc.add_heading('Dietary Assessment and Food Groups', level=2)
    p = doc.add_paragraph(
        'Dietary intake was assessed using a validated semi-quantitative food frequency questionnaire. Foods were '
        'aggregated into 12 groups: grain products, protein foods, vegetables, dairy products, fruits, fried foods, '
        'high-fat meat, processed foods, sugar-sweetened beverages, additional salt use, salty foods, and sweet foods. '
        'Each group was scored on 3- or 4-point scales based on consumption frequency and adherence to Korean dietary '
        'guidelines (18). For network construction, scores were binarized (high consumption: score ≥3; low consumption: '
        'score <3), representing adequate intake for healthy foods and frequent consumption for unhealthy foods.'
    )
    
    doc.add_heading('Network Construction and Analysis', level=2)
    p = doc.add_paragraph(
        'For each stratified group, we constructed co-occurrence networks representing simultaneous high-consumption '
        'patterns. Co-occurrence frequency between food groups i and j was calculated as the proportion of participants '
        'consuming both at high levels. Edges were retained if co-occurrence exceeded the 70th percentile within each '
        'group, yielding undirected weighted networks. We calculated three node centrality metrics: degree centrality '
        '(number of direct connections), betweenness centrality (frequency on shortest paths between other nodes), and '
        'closeness centrality (inverse average shortest path length to other nodes). Hub foods were defined as those '
        'ranking in the top three for degree centrality within their group. All analyses used Python 3.9 with NetworkX 2.8.'
    )
    
    doc.add_page_break()
    
    # Results
    doc.add_heading('Results', level=1)
    
    doc.add_heading('Network Structure Consistency', level=2)
    p = doc.add_paragraph(
        'All 11 stratified networks exhibited identical topology with 12 nodes, 20 edges, and density of 0.303 '
        '(30.3% of possible connections present). This structural consistency reflects the co-occurrence threshold '
        'methodology, which adapted to group-specific consumption patterns while maintaining comparable network density. '
        'Despite this topological uniformity, centrality distributions varied substantially across groups, indicating '
        'that identical structures can harbor different functional organizations.'
    )
    
    doc.add_heading('Universal Hub Foods', level=2)
    p = doc.add_paragraph(
        'Three food groups consistently emerged as hubs (top-3 degree centrality) across all 11 networks: protein foods, '
        'vegetables, and grain products. Protein foods exhibited the highest centrality (degree: 0.636-1.000; median: 0.818), '
        'appearing as the top-ranked hub in 10 of 11 groups. Vegetables showed similarly high centrality (degree: 0.455-1.000; '
        'median: 0.727) and ranked first or second in all groups. Grain products displayed moderate but universal centrality '
        '(degree: 0.364-0.545; median: 0.455), consistently ranking among the top three hubs. These three foods collectively '
        'represent a "universal dietary core" transcending demographic and metabolic boundaries.'
    )
    
    doc.add_heading('Age-Related Hub Transitions', level=2)
    p = doc.add_paragraph(
        'Hub composition exhibited systematic age-related shifts, particularly pronounced in males. Among young adults '
        '(19-39 years), sugar-sweetened beverages appeared as a top-3 hub in both males with and without MetS (degree: '
        '0.273-0.364). This beverage centrality diminished in middle age (40-59 years), replaced by grain products as '
        'a more prominent hub. In older adults (60-74 years), grain products showed elevated centrality (degree: 0.455-0.545), '
        'particularly in males. Among females, sweet food consumption emerged as a hub in young adults but not in older groups, '
        'suggesting sex-specific dietary transitions across the lifespan.'
    )
    
    doc.add_heading('Sex Differences in Hub Patterns', level=2)
    p = doc.add_paragraph(
        'Female networks showed higher centrality for vegetables compared to male networks across age and MetS groups. '
        'Additionally, sweet food consumption appeared as a hub in young females (degree: 0.364) but not in age-matched '
        'males, where sugar-sweetened beverages were more central. These patterns suggest sex-differentiated dietary '
        'preferences that persist across metabolic health status.'
    )
    
    doc.add_heading('Metabolic Syndrome Effects', level=2)
    p = doc.add_paragraph(
        'MetS status showed limited influence on hub food identity, with protein foods, vegetables, and grain products '
        'maintaining hub status across MetS+ and MetS- groups within the same age-sex stratum. However, subtle differences '
        'emerged in hub rankings: in middle-aged and older females with MetS+, vegetables occasionally ranked first, displacing '
        'protein foods to second position. This may reflect dietary modification following MetS diagnosis or health-conscious '
        'eating patterns among individuals with metabolic abnormalities.'
    )
    
    doc.add_page_break()
    
    # Discussion
    doc.add_heading('Discussion', level=1)
    
    doc.add_heading('Principal Findings', level=2)
    p = doc.add_paragraph(
        'This stratified network analysis of 22,964 Korean adults revealed both universal and group-specific patterns in '
        'dietary hubs. Despite uniform network topology across 11 demographic-metabolic subgroups, centrality distributions '
        'varied substantially. Three foods—protein foods, vegetables, and grain products—emerged as universal hubs suitable '
        'for population-wide intervention targets. Concurrently, age-related hub transitions (sugar-sweetened beverages to '
        'grain products) and sex differences (sweet foods in females vs. sugar-sweetened beverages in males) support the need '
        'for personalized dietary counseling. These findings advance dietary pattern research by demonstrating that network '
        'analysis can simultaneously identify both common targets and tailored strategies.'
    )
    
    doc.add_heading('Comparison with Previous Research', level=2)
    p = doc.add_paragraph(
        'Previous dietary network studies have identified hub foods in overall populations, including protein sources, vegetables, '
        'and grains (9, 19, 20), consistent with our universal hubs. However, by stratifying analyses, we demonstrate that hub '
        'prominence varies by demographic characteristics. Our finding that sugar-sweetened beverage centrality declines with age '
        'aligns with epidemiological data showing age-related reductions in beverage consumption (21, 22), but extends this '
        'observation to the network structural level. Similarly, sex differences in sweet food consumption have been documented '
        'through traditional dietary assessment (23, 24), but our network approach reveals how these preferences integrate into '
        'broader dietary patterns differently for males and females.'
    )
    
    doc.add_heading('Clinical and Public Health Implications', level=2)
    p = doc.add_paragraph(
        'Our findings support a two-tiered approach to dietary intervention: (1) promote the protein-vegetable-grain triad as '
        'universal dietary foundation applicable across all demographic groups; (2) tailor additional recommendations to age and '
        'sex, such as emphasizing reduced sugar-sweetened beverage intake in young adults or addressing age-appropriate grain '
        'consumption patterns in older adults. For MetS management, the universal hubs provide actionable targets that require '
        'minimal customization, potentially improving intervention uptake and sustainability. Healthcare providers can counsel '
        'patients to build meals around protein, vegetables, and grains, confident that this core recommendation applies broadly '
        'while allowing room for individual preferences.'
    )
    
    doc.add_heading('Methodological Considerations', level=2)
    p = doc.add_paragraph(
        'We chose co-occurrence network analysis for its interpretability and direct representation of simultaneous consumption '
        'patterns, making results accessible to clinicians and policymakers. Alternative approaches such as Gaussian graphical '
        'models estimate partial correlations controlling for all other foods but require larger samples and impose stronger '
        'statistical assumptions. Our consistent network topology across groups resulted from the percentile-based thresholding '
        'methodology, which prioritizes internal consistency within groups rather than absolute co-occurrence frequencies. While '
        'this approach facilitates between-group comparisons of centrality distributions, it precludes direct comparison of '
        'absolute edge weights across networks.'
    )
    
    doc.add_heading('Strengths and Limitations', level=2)
    p = doc.add_paragraph(
        'Strengths include the large, nationally representative sample, comprehensive stratification strategy, and multiple '
        'centrality metrics providing convergent evidence for hub identification. Limitations merit consideration. First, the '
        'cross-sectional design precludes causal inference regarding diet-MetS relationships. Second, food frequency questionnaires '
        'are subject to recall bias, though validation studies support their use in ranking dietary intake. Third, our 12-food '
        'group classification, while based on Korean dietary patterns, may not capture all relevant granularity. Fourth, the '
        'female young adult MetS+ group was excluded due to small sample size, limiting generalizability to this stratum. Finally, '
        'our binarization of dietary scores (≥3 vs. <3) simplified the analysis but discarded information about consumption intensity.'
    )
    
    doc.add_heading('Conclusions', level=2)
    p = doc.add_paragraph(
        'Stratified network analysis reveals that protein foods, vegetables, and grain products constitute a universal dietary '
        'core across sex, age, and metabolic health subgroups in Korean adults, providing robust targets for population-wide '
        'interventions. Simultaneously, age-specific and sex-specific hub patterns support personalized nutrition strategies. '
        'Future research should extend these methods to longitudinal data to examine temporal stability of network structures '
        'and validate hub-focused interventions in randomized trials. The integration of network analysis with traditional '
        'nutritional epidemiology offers a promising framework for developing both universal and tailored dietary recommendations.'
    )
    
    doc.add_page_break()
    
    # References (placeholder)
    doc.add_heading('References', level=1)
    p = doc.add_paragraph(
        '1. Alberti KG, Eckel RH, Grundy SM, et al. Harmonizing the metabolic syndrome. Circulation 2009;120:1640-5.\n\n'
        '2. Saklayen MG. The Global Epidemic of the Metabolic Syndrome. Curr Hypertens Rep 2018;20:12.\n\n'
        '3. Kastorini CM, Milionis HJ, Esposito K, et al. The effect of Mediterranean diet on metabolic syndrome and its '
        'components: a meta-analysis. J Am Coll Cardiol 2011;57:1299-313.\n\n'
        '4. Dietary Guidelines Advisory Committee. Scientific Report of the 2020 Dietary Guidelines Advisory Committee. '
        'Washington, DC: U.S. Department of Agriculture, 2020.\n\n'
        '[Additional references to be added]'
    )
    
    return doc

def main():
    """메인 함수"""
    print("Creating SCI-level English paper...")
    
    doc = create_sci_paper()
    
    # 저장
    output_dir = Path(__file__).parent.parent / 'result' / 'manuscript'
    output_file = output_dir / 'Paper2_SCI_English.docx'
    
    doc.save(output_file)
    
    print(f"\n✓ SCI paper saved: {output_file}")
    print(f"  File size: {output_file.stat().st_size / 1024:.1f} KB")
    print("\n영어 SCI급 논문 작성 완료!")

if __name__ == '__main__':
    main()
