#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Enhanced Korean Report Generator for Paper 2
Includes age/sex-specific coaching strategies and strengthened MetS impact analysis
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_korean_report():
    """Create enhanced Korean report with coaching strategies and MetS analysis"""
    
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('한국 성인 22,964명의 식이 동시섭취 네트워크에서 나타나는\n보편적 Hub 식품과 집단 특이적 Hub 식품:\n층화 분석 연구')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # =================================================================
    # 1. 연구 개요
    # =================================================================
    doc.add_heading('1. 연구 개요', level=1)
    
    doc.add_heading('가. 연구 배경 및 필요성', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('기존 식이 패턴 분석의 한계')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('전통적 접근법(요인분석, 군집분석): 전체적 패턴만 파악')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('개별 식품 간 동시섭취 관계를 명확히 규명하기 어려움')
    
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('네트워크 분석의 강점')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('식품 간 연결 관계를 시각화하고 정량화')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('중심 식품(hub)을 객관적으로 식별 가능')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('맞춤형 영양 중재 전략 개발에 유용')
    
    doc.add_heading('나. 연구 목적', level=2)
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('보편적 Hub 식품 식별')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('성별, 연령, 대사증후군 상태와 무관하게 일관되게 나타나는 중심 식품 규명')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('집단 특이적 Hub 식품 탐색')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('인구집단별로 특징적으로 나타나는 Hub 식품 파악')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('대사증후군 영향 분석')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('MetS 유무에 따른 Hub 식품 구성 및 순위 변화 확인')
    
    doc.add_heading('다. 연구 대상', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('데이터: ')
    run.bold = True
    p.add_run('국민건강영양조사 2019-2021 (3개년 통합)')
    
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('대상: ')
    run.bold = True
    p.add_run('19-74세 성인 22,964명')
    
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('층화 기준: ')
    run.bold = True
    p.add_run('성별(2) × 연령대(3) × 대사증후군(2) = 11개 집단')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('연령대: 청년층(19-39세), 중년층(40-59세), 장년층(60-74세)')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('대사증후군: NCEP ATP III 기준 (복부비만은 아시아 기준 적용)')
    
    # =================================================================
    # 2. 연구 방법
    # =================================================================
    doc.add_heading('2. 연구 방법', level=1)
    
    doc.add_heading('가. 식품 분류', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('13개 식품군 분류')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('곡류, 단백질 식품, 채소, 과일, 유제품, 당류 첨가 음료, 기호식품(커피/차/주류), 달콤한 음식, 패스트푸드, 라면, 소스/조미료, 간식, 김치류')
    
    doc.add_heading('나. 섭취 빈도 점수화', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('5점 척도 변환:')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('1점(월 1회 미만) → 5점(거의 매일)')
    
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('이진화 기준: ')
    run.bold = True
    p.add_run('점수 ≥3 → 적정/빈번 섭취로 정의')
    
    doc.add_heading('다. 네트워크 구성', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('동시섭취(co-occurrence) 네트워크')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('두 식품군을 모두 점수 ≥3으로 섭취하는 사람 비율로 연결 강도 산출')
    
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('중심성 지표 계산:')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('연결 중심성(Degree centrality): 직접 연결 수')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('매개 중심성(Betweenness centrality): 경로 중개 역할')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('근접 중심성(Closeness centrality): 전체 식품과의 근접도')
    
    doc.add_heading('라. Hub 식품 정의', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('각 층화 집단(11개)별 Top-3 연결 중심성 식품을 Hub로 정의')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('보편적 Hub: ')
    run.bold = True
    p.add_run('11개 집단 모두에서 Top-3에 포함되는 식품')
    
    # =================================================================
    # 3. 주요 결과
    # =================================================================
    doc.add_heading('3. 주요 결과', level=1)
    
    doc.add_heading('가. 보편적 Hub 식품', level=2)
    p = doc.add_paragraph()
    run = p.add_run('■ 3개 식품군이 모든 집단에서 일관되게 Top-3 Hub로 확인:')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('단백질 식품(Protein Foods)')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('육류, 생선, 해산물, 달걀, 콩류 포함')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('한국인 식사에서 중심적 역할')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('채소(Vegetables)')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('다양한 채소류 (김치 제외)')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('식사 구성의 필수 요소')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('곡류(Grain Products)')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('밥, 빵, 면류 등 주식')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('한국인 식사 구조의 핵심')
    
    p = doc.add_paragraph()
    run = p.add_run('→ 의의: 성별, 연령, 건강 상태와 무관한 "식사 구조의 뼈대"')
    run.italic = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_heading('나. 집단 특이적 Hub 식품', level=2)
    
    # 남성 청년층
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('남성 청년층(19-39세)')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('특징: 당류 첨가 음료(Sugar-Sweetened Beverages)가 3위 Hub')
    run.font.color.rgb = RGBColor(204, 0, 0)
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('MetS 유무와 관계없이 일관된 패턴')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('해석: 젊은 남성의 음료 섭취 습관 반영')
    
    # 여성 청년층
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('여성 청년층(19-39세)')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('특징: 달콤한 음식(Sweet Food)이 3위 Hub (MetS(-) 집단)')
    run.font.color.rgb = RGBColor(204, 0, 0)
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('MetS(+) 집단은 없어 비교 불가 (표본수 부족)')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('해석: 젊은 여성의 간식/디저트 섭취 패턴')
    
    # 중년층/장년층
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('중·장년층(40세 이상)')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('특징: 곡류가 3위 Hub로 고정')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('예외: 남성 장년층 MetS(+)에서 곡류가 2위로 상승')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('해석: 나이 들수록 전통적 주식 중심 식사 패턴 강화')
    
    doc.add_heading('다. Hub 순위 분석 (Top-3 내 순서)', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('■ 대부분 집단에서 "단백질-채소-곡류" 순서 유지')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('예외 사례:')
    run.bold = True
    
    # Exception 1
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('여성 장년층 MetS(+): ')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    run = p.add_run('채소-곡류-단백질')
    run.font.color.rgb = RGBColor(204, 0, 0)
    p = doc.add_paragraph('', style='List Bullet 3')
    p.add_run('채소가 1위로 부상! (가장 드라마틱한 변화)')
    
    # Exception 2
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('남성 장년층 MetS(+): ')
    run.bold = True
    run = p.add_run('단백질-곡류-채소')
    p = doc.add_paragraph('', style='List Bullet 3')
    p.add_run('곡류가 2위로 상승')
    
    # Exception 3
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('여성 장년층 MetS(-): ')
    run.bold = True
    run = p.add_run('단백질-곡류-채소')
    p = doc.add_paragraph('', style='List Bullet 3')
    p.add_run('곡류가 채소보다 우선')
    
    doc.add_heading('라. 대사증후군의 영향 ★★★ 중요', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('■ 핵심 발견: Hub 식품 구성에는 제한적 영향')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    run.font.size = Pt(12)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('Hub 구성(어떤 식품이 Hub인가)은 대체로 일관됨')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('단백질-채소-곡류는 MetS 유무와 관계없이 hub 유지')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('청년층의 특이적 hub(음료/단 음식)도 MetS와 무관')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('Hub 순위(1위, 2위, 3위 순서)에는 미세한 변화 관찰')
    run.bold = True
    
    # 여성 장년층 상세 분석
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('【여성 장년층】 가장 두드러진 차이:')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph('', style='List Bullet 3')
    run = p.add_run('MetS(+): 채소-곡류-단백질')
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    # Use plain paragraph with manual indentation for deeper nesting
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.2)
    run = p.add_run('→ 채소가 1위로 부상! (단백질은 3위로 하락)')
    run.italic = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph('', style='List Bullet 3')
    run = p.add_run('MetS(-): 단백질-곡류-채소')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.2)
    run = p.add_run('→ 전형적인 순서 유지')
    run.italic = True
    
    # 남성 장년층 분석
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('【남성 장년층】')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet 3')
    run = p.add_run('MetS(+): 단백질-곡류-채소')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.2)
    run = p.add_run('→ 곡류가 2위로 상승 (채소는 3위로 하락)')
    run.italic = True
    
    p = doc.add_paragraph('', style='List Bullet 3')
    run = p.add_run('MetS(-): 단백질-채소-곡류')
    
    # 기타 집단
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('【기타 집단】')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 3')
    p.add_run('청년층, 중년층: MetS 유무에 따른 순위 변화 미미')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('해석')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('장년층에서만 MetS 영향이 뚜렷 → 나이와의 상호작용 시사')
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('여성 장년층 MetS(+)의 채소 우선 패턴:')
    run.bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.2)
    p.add_run('• 진단 후 식이 변화(건강관리 의식) 반영 가능성')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.2)
    p.add_run('• 근감소증 예방을 위한 채소 섭취 강조 교육의 결과일 수 있음')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.2)
    p.add_run('• 그러나 단백질이 3위로 밀려난 점은 우려 → 영양 코칭 필요')
    
    # =================================================================
    # 4. 인구집단별 맞춤형 영양 코칭 전략 ★★★ NEW SECTION
    # =================================================================
    doc.add_heading('4. 인구집단별 맞춤형 영양 코칭 전략', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('■ 연구 결과를 바탕으로 한 실용적 영양 중재 방안')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(12)
    
    doc.add_heading('가. 남성 코칭 전략', level=2)
    
    # 남성 청년층
    p = doc.add_paragraph()
    run = p.add_run('【남성 청년층 19-39세】')
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('핵심 이슈: 당류 첨가 음료가 Hub (3위)')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('MetS 유무와 관계없이 일관된 패턴 → 연령층 전체의 습관')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('코칭 전략')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('음료 섭취 습관 개선을 최우선 목표로 설정')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 당류 첨가 음료 → 물, 탄산수, 무가당 음료로 대체')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• "탄산음료 1캔 = 각설탕 10개" 등 구체적 정보 제공')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('단백질-채소 중심 식사는 유지하면서 음료만 교체')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 현재 식사 구조를 크게 바꾸지 않아 실천 용이')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS 예방 관점:')
    run.bold = True
    p.add_run(' 음료 개선만으로도 당 섭취량 대폭 감소 효과')
    
    # 남성 중년층
    p = doc.add_paragraph()
    run = p.add_run('【남성 중년층 40-59세】')
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('특징: 단백질-채소-곡류 순서 (MetS 유무 관계없이)')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('코칭 전략')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(-) 집단:')
    run.bold = True
    p.add_run(' 현재 패턴 유지 + 예방 차원 관리')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 단백질-채소 중심은 바람직한 구조')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 곡류의 질 개선: 현미, 잡곡밥 비율 높이기')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 정기 건강검진으로 MetS 조기 발견')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(+) 집단:')
    run.bold = True
    p.add_run(' 곡류 질 개선 + 부분 조절')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 흰 쌀밥 → 현미/잡곡밥 단계적 전환')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 곡류 섭취량 약간 줄이고 채소 비중 증가')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 단백질은 지방 적은 부위 선택(살코기, 생선, 콩류 증가)')
    
    # 남성 장년층
    p = doc.add_paragraph()
    run = p.add_run('【남성 장년층 60-74세】')
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('핵심 차이: MetS(+)에서 곡류가 2위로 상승')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('MetS(+): 단백질-곡류-채소')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('MetS(-): 단백질-채소-곡류')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('코칭 전략')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(-) 집단:')
    run.bold = True
    p.add_run(' 현재 패턴 유지 + 단백질 강화')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 근감소증 예방을 위한 단백질 섭취 지속')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 채소 섭취 유지로 만성질환 예방')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(+) 집단:')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    p.add_run(' 곡류 질 개선 + 채소 비중 증가 필요')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    run = p.add_run('• 문제: 곡류가 과도하게 높은 순위 (2위)')
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 백미 의존도를 낮추고 통곡물로 전환')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    run = p.add_run('• 채소가 3위로 밀린 점 개선: 채소 반찬 가짓수 늘리기')
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 단백질은 유지하되 저지방 선택 (생선, 두부, 닭가슴살)')
    
    doc.add_heading('나. 여성 코칭 전략', level=2)
    
    # 여성 청년층
    p = doc.add_paragraph()
    run = p.add_run('【여성 청년층 19-39세】')
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('핵심 이슈: 달콤한 음식이 Hub (3위)')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('MetS(+) 집단 데이터 부족 (유병률 낮음)')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('코칭 전략')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('간식/디저트 섭취 패턴 개선')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 달콤한 음식 → 과일, 견과류, 요거트로 대체')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 섭취 빈도 줄이기: 매일 → 주 2-3회')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('단백질-채소 중심 식사는 유지')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 현재 식사 구조는 바람직함')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS 예방 관점:')
    run.bold = True
    p.add_run(' 당 섭취 조절 + 임신·출산 대비 영양 관리')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 철분, 엽산 등 미량영양소 충분 섭취')
    
    # 여성 중년층
    p = doc.add_paragraph()
    run = p.add_run('【여성 중년층 40-59세】')
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('특징: 단백질-채소-곡류 순서 (MetS 유무 관계없이)')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('코칭 전략')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(-) 집단:')
    run.bold = True
    p.add_run(' 현재 패턴 유지 + 폐경기 대비')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 균형 잡힌 식사 패턴 지속')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 칼슘(유제품), 비타민 D 강화 (골다공증 예방)')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 정기 건강검진 (폐경 후 MetS 위험 증가 대비)')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(+) 집단:')
    run.bold = True
    p.add_run(' 곡류 질 개선 + 체중 관리')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 정제 곡류 줄이고 통곡물 증가')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 복부비만 관리: 전체 섭취량 조절 + 운동')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 단백질 유지하되 식물성 단백질(콩, 두부) 비중 높이기')
    
    # 여성 장년층 ★★★ 가장 중요
    p = doc.add_paragraph()
    run = p.add_run('【여성 장년층 60-74세】 ★★★ 가장 중요')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph()
    run = p.add_run('※ 대사증후군 영향이 가장 뚜렷하게 나타나는 집단!')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    run.font.size = Pt(11)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('Hub 순위의 극적 역전 현상')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(+): 채소-곡류-단백질')
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    run = p.add_run('→ 채소가 1위로 부상! (단백질은 3위로 하락)')
    run.italic = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('MetS(-): 단백질-곡류-채소')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('→ 전형적인 순서')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('해석')
    run.bold = True
    
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('긍정적 측면: 채소 섭취 증가 (건강 의식 반영)')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• MetS 진단 후 채소 중심 식단으로 변화 노력')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 혈당, 혈압 관리를 위한 채소 강조 교육 효과')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('우려 사항: 단백질이 3위로 하락')
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    run = p.add_run('• 장년층에서 단백질 부족은 근감소증 위험 증가')
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 채소 강조에 집중하다 단백질 소홀히 한 가능성')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('코칭 전략 (최우선 타겟 집단)')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(-) 집단:')
    run.bold = True
    p.add_run(' 곡류 질 개선 + 채소 증가')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 단백질 유지하며 채소 섭취 늘리기')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 곡류를 통곡물로 전환')
    
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('MetS(+) 집단:')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    p.add_run(' 단백질 재강화 필수!')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    run = p.add_run('• 핵심 메시지: "채소는 좋지만, 단백질도 필수!"')
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 매 끼니 단백질 반찬 1가지 이상 포함')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('  예: 생선구이, 두부조림, 달걀찜, 콩나물무침')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 저지방 단백질 우선: 생선, 닭가슴살, 두부, 콩류')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 채소는 현재 수준 유지 (이미 충분)')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.add_run('• 곡류는 통곡물로 전환 + 양 조절')
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    run = p.add_run('• 근감소증 예방 교육: 단백질의 중요성 강조')
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    doc.add_heading('다. 코칭 전략 요약 및 우선순위', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('■ 중재 우선순위')
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('1순위: 여성 장년층 MetS(+)')
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('단백질 재강화 (근감소증 예방)')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('2순위: 남성 장년층 MetS(+)')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('곡류 질 개선 + 채소 비중 증가')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('3순위: 남성 청년층 (전체)')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('당류 첨가 음료 섭취 감소')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('4순위: 여성 청년층')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('달콤한 음식 섭취 조절')
    
    p = doc.add_paragraph()
    run = p.add_run('■ 공통 전략')
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('보편적 Hub(단백질-채소-곡류) 유지를 기본으로 함')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('집단별 특이적 문제(음료, 단 음식, 순위 불균형)에 집중')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('MetS(+) 집단: 곡류의 질 개선 + 전체 섭취량 조절')
    
    # =================================================================
    # 5. 결론 및 제언
    # =================================================================
    doc.add_heading('5. 결론 및 제언', level=1)
    
    doc.add_heading('가. 주요 결론', level=2)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('보편적 Hub 식품 확인')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('단백질-채소-곡류는 성별, 연령, MetS 상태와 무관한 핵심 식사 구조')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('한국인의 기본 식사 뼈대로서 유지 필요')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('집단별 특성 반영 필요')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('청년층: 음료/간식 개선이 핵심')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('중·장년층: 곡류 질 개선 및 균형 유지')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('대사증후군 영향은 제한적이나 장년층에서 순위 변화 관찰')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    run = p.add_run('여성 장년층 MetS(+): 채소 1위, 단백질 3위 → 단백질 재강화 필요')
    run.font.color.rgb = RGBColor(204, 0, 0)
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('남성 장년층 MetS(+): 곡류 2위 → 질 개선 필요')
    
    doc.add_heading('나. 정책 제언', level=2)
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('맞춤형 영양 교육 프로그램 개발')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('연령·성별·건강상태별 차별화된 교육 콘텐츠')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('특히 여성 장년층 MetS(+) 집단 대상 단백질 강화 교육')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('청년층 대상 음료/간식 섭취 개선 캠페인')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('직장, 대학교 등에서 실행 가능한 환경 조성')
    
    p = doc.add_paragraph('', style='List Number')
    run = p.add_run('장년층 대상 근감소증 예방 영양 관리')
    run.bold = True
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('단백질 섭취의 중요성에 대한 인식 제고')
    p = doc.add_paragraph('', style='List Bullet 2')
    p.add_run('저지방 단백질 식품 접근성 향상')
    
    doc.add_heading('다. 연구의 의의', level=2)
    
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('방법론적 의의: ')
    run.bold = True
    p.add_run('네트워크 분석을 통한 식이 패턴의 새로운 접근')
    
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('실용적 의의: ')
    run.bold = True
    p.add_run('인구집단별 맞춤형 영양 중재 전략 제시')
    
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('정책적 의의: ')
    run.bold = True
    p.add_run('국가 영양 정책 수립의 과학적 근거 제공')
    
    doc.add_heading('라. 연구의 제한점', level=2)
    
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('횡단연구 설계로 인과관계 추론 제한')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('자가보고식 식품섭취빈도조사의 한계 (회상 오류)')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('일부 집단(여성 청년층 MetS(+))은 표본수 부족으로 분석 불가')
    
    doc.add_heading('마. 향후 연구 방향', level=2)
    
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('종단연구를 통한 식이 패턴 변화 추적')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Hub 식품 기반 영양 중재의 효과 검증')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('다른 국가/문화권과의 비교 연구')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('더 세분화된 층화 분석 (소득, 교육 수준 등)')
    
    # Save document
    output_path = '/home/user/webapp/ver3.0_2511/result/manuscript/Paper2_Korean_Report.docx'
    doc.save(output_path)
    print(f"✓ Korean report saved: {output_path}")
    return output_path

if __name__ == "__main__":
    print("Creating enhanced Korean report with coaching strategies...")
    output_path = create_korean_report()
    print(f"\n✓ Complete! File saved to: {output_path}")
    print("\nKey enhancements:")
    print("  - Section 4: Population-specific coaching strategies")
    print("  - Enhanced Section 3-라: Strengthened MetS impact analysis")
    print("  - Fixed Word style errors (max 3 bullet levels)")
    print("  - Emphasis on female elderly MetS(+) protein reinforcement")
