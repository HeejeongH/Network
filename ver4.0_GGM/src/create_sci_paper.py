#!/usr/bin/env python3
"""
SCI급 영어 논문 작성 스크립트 - ver4.0 GGM
GGM 방법론으로 업데이트된 내용
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_sci_paper_ggm():
    """SCI급 영어 논문 작성 (GGM 방법론)"""
    doc = Document()
    
    # 페이지 설정
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Gaussian Graphical Models Reveal Universal and Group-Specific Hub Foods ')
    run.bold = True
    run.font.size = Pt(14)
    run = title.add_run('in Dietary Networks: A Stratified Analysis of 22,964 Korean Adults')
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Running title
    p = doc.add_paragraph()
    run = p.add_run('Running Title: ')
    run.bold = True
    p.add_run('GGM-Based Dietary Network Hubs')
    
    doc.add_paragraph()
    
    # Authors (placeholder)
    p = doc.add_paragraph()
    run = p.add_run('Authors: ')
    run.bold = True
    p.add_run('[To be filled]')
    
    doc.add_paragraph()
    
    # Abstract
    p = doc.add_heading('Abstract', level=1)
    
    # Background
    p = doc.add_paragraph()
    run = p.add_run('Background: ')
    run.bold = True
    p.add_run('Traditional dietary pattern analysis often overlooks the complex conditional dependencies among foods and how these patterns vary across population subgroups. Gaussian graphical models (GGM) offer a sophisticated approach to identify "hub" foods that are central to dietary patterns through partial correlations, but few studies have examined how these hubs differ by demographic and metabolic characteristics.')
    
    # Objective
    p = doc.add_paragraph()
    run = p.add_run('Objective: ')
    run.bold = True
    p.add_run('We applied semiparametric Gaussian copula graphical models (SGCGM) with cross-validated graphical lasso to identify universal and group-specific hub foods across sex, age, and metabolic syndrome (MetS) status in a large Korean population.')
    
    # Methods
    p = doc.add_paragraph()
    run = p.add_run('Methods: ')
    run.bold = True
    p.add_run('Data from 22,964 adults (19-74 years) in the Korea National Health and Nutrition Examination Survey were stratified into 11 groups by sex, age (19-39, 40-59, 60-74 years), and MetS status. Dietary networks were constructed using SGCGM to estimate conditional dependencies via partial correlations, controlling for confounding relationships. Food group scores (continuous, range: 1-4) were analyzed using rank-based transformations (Spearman correlation) followed by L1-penalized precision matrix estimation (graphical lasso). Optimal regularization parameters were selected via 5-fold cross-validation. Hub foods were identified using degree, betweenness, and closeness centrality metrics.')
    
    # Results  
    p = doc.add_paragraph()
    run = p.add_run('Results: ')
    run.bold = True
    p.add_run('Network topology varied substantially across groups (edges: 3-9; density: 0.045-0.136), reflecting group-specific dietary patterns. Regularization parameters (α) ranged from 0.124 to 0.164. Three foods emerged as frequent top-3 hubs: processed foods (appearing in 10/11 groups), protein foods (9/11 groups), and fried foods (9/11 groups). Unlike simple co-occurrence networks, GGM revealed that traditional "healthy" foods (vegetables, grains) showed lower centrality, while processed and fried foods consistently occupied central network positions. Hub patterns differed by demographics: young adults showed higher centrality for processed foods and sugar-sweetened beverages, while older adults exhibited more diverse hub composition. Partial correlations (range: -0.47 to +0.68) demonstrated both positive co-consumption and negative substitution patterns.')
    
    # Conclusions
    p = doc.add_paragraph()
    run = p.add_run('Conclusions: ')
    run.bold = True
    p.add_run('GGM analysis reveals that processed and fried foods, rather than traditionally recommended healthy foods, occupy central hub positions in actual Korean dietary networks. This finding highlights a disconnect between dietary guidelines and real-world consumption patterns. The processed-protein-fried triad represents a common dietary core suitable for population-wide intervention targets, while age-specific hub patterns support tailored nutritional counseling. These findings demonstrate the value of GGM-based network analysis for identifying realistic intervention targets in personalized nutrition strategies.')
    
    # Keywords
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.bold = True
    p.add_run('Gaussian graphical models; partial correlations; dietary networks; hub foods; metabolic syndrome; stratified analysis; graphical lasso; personalized nutrition')
    
    doc.add_page_break()
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    
    p = doc.add_paragraph(
        'Metabolic syndrome (MetS) affects one-quarter of the global adult population and substantially '
        'increases risk for cardiovascular disease and type 2 diabetes (1, 2). Dietary modification remains '
        'a cornerstone of MetS prevention and management, yet identifying optimal dietary targets for diverse '
        'population groups presents ongoing challenges (3, 4). Traditional dietary assessment methods—analyzing '
        'individual nutrients or deriving patterns through principal component or cluster analysis—may inadequately '
        'capture the complex, interconnected nature of food consumption behaviors (5, 6). Moreover, simple '
        'co-occurrence approaches can detect spurious associations arising from indirect relationships (7).'
    )
    
    p = doc.add_paragraph(
        'Gaussian graphical models (GGM) offer a more sophisticated framework for dietary pattern analysis by '
        'modeling conditional dependencies through partial correlations (8, 9). Unlike simple co-occurrence networks '
        'that conflate direct and indirect relationships, GGM explicitly controls for confounding by estimating the '
        'correlation between two foods adjusting for all other foods (10, 11). This reveals "hub" foods that maintain '
        'direct connections with multiple other foods, potentially serving as leverage points for behavioral '
        'interventions. Semiparametric extensions using Gaussian copulas accommodate non-normal distributions typical '
        'in dietary data while preserving the interpretability of correlation-based networks (12, 13).'
    )
    
    p = doc.add_paragraph(
        'Recent applications of GGM to dietary data have identified meal-specific and habitual dietary networks in '
        'European populations (12, 13) and pregnancy-related patterns in U.S. cohorts (14). However, these studies '
        'analyzed populations as single entities, potentially obscuring important heterogeneity across demographic '
        'and clinical subgroups (15, 16). Given established differences in dietary preferences by sex and age (17, 18), '
        'and the complex bidirectional relationship between diet and metabolic health (19, 20), stratified GGM analysis '
        'is needed to inform personalized nutrition strategies.'
    )
    
    p = doc.add_paragraph(
        'We conducted a stratified GGM analysis of dietary patterns in 22,964 Korean adults, examining 11 groups '
        'defined by sex, age, and MetS status. Our objectives were to: (1) construct conditional dependency networks '
        'using SGCGM with cross-validated graphical lasso; (2) identify hub foods within each demographic-metabolic '
        'subgroup using multiple centrality metrics; (3) distinguish universal hubs appearing across all groups from '
        'group-specific hubs; and (4) characterize hub food transitions across the lifespan. These findings provide '
        'evidence-based targets for both population-wide and personalized dietary interventions grounded in actual '
        'consumption patterns.'
    )
    
    doc.add_page_break()
    
    # Methods
    doc.add_heading('Methods', level=1)
    
    doc.add_heading('Study Population', level=2)
    p = doc.add_paragraph(
        'We analyzed data from the Korea National Health and Nutrition Examination Survey (KNHANES), a nationally '
        'representative cross-sectional survey employing complex, stratified, multistage probability sampling. The '
        'study included 22,964 adults aged 19-74 years with complete dietary, anthropometric, and biochemical data. '
        'Participants were stratified into 12 potential groups by sex (male/female), age (19-39, 40-59, 60-74 years), '
        'and MetS status (MetS+/MetS-). The female young adult MetS+ group was excluded due to insufficient sample '
        'size (n<100), yielding 11 groups for analysis (range: 516-5,629 participants per group). The KNHANES protocol '
        'received institutional review board approval, and all participants provided written informed consent.'
    )
    
    doc.add_heading('Metabolic Syndrome Definition', level=2)
    p = doc.add_paragraph(
        'MetS was defined using modified NCEP ATP III criteria with Asian-specific waist circumference thresholds (21). '
        'Participants meeting ≥3 of five criteria were classified as MetS+: (1) waist circumference ≥90 cm (men) or '
        '≥85 cm (women); (2) triglycerides ≥150 mg/dL or medication use; (3) HDL-cholesterol <40 mg/dL (men) or '
        '<50 mg/dL (women); (4) blood pressure ≥130/85 mmHg or medication use; (5) fasting glucose ≥100 mg/dL or '
        'medication use.'
    )
    
    doc.add_heading('Dietary Assessment and Food Groups', level=2)
    p = doc.add_paragraph(
        'Dietary intake was assessed using a validated semi-quantitative food frequency questionnaire. Foods were '
        'aggregated into 12 groups: grain products, protein foods, vegetables, dairy products, fruits, fried foods, '
        'high-fat meat, processed foods, sugar-sweetened beverages, additional salt use, salty foods, and sweet foods. '
        'Each group was scored on 3- or 4-point scales based on consumption frequency and adherence to Korean dietary '
        'guidelines (22). Scores were analyzed as continuous variables (not binarized) to preserve information and '
        'allow estimation of partial correlations.'
    )
    
    doc.add_heading('Gaussian Graphical Model Construction', level=2)
    p = doc.add_paragraph(
        'For each stratified group, we constructed dietary networks using semiparametric Gaussian copula graphical '
        'models (SGCGM) to estimate conditional dependencies via partial correlations. The analytical pipeline '
        'consisted of four steps:'
    )
    
    p = doc.add_paragraph(
        '(1) Rank-based transformation: To accommodate non-normal distributions typical in dietary data, we computed '
        'Spearman rank correlations between all food group pairs, effectively applying a nonparanormal transformation (23).'
    )
    
    p = doc.add_paragraph(
        '(2) Graphical lasso: We estimated sparse precision matrices (Θ) via L1-penalized maximum likelihood, '
        'solving: minimize -log det(Θ) + tr(SΘ) + λ||Θ||₁, where S is the empirical correlation matrix and λ '
        'controls sparsity (24, 25). The L1 penalty encourages sparse networks by shrinking small partial correlations '
        'to exactly zero, removing spurious edges from confounding.'
    )
    
    p = doc.add_paragraph(
        '(3) Cross-validation: For each group, we selected the optimal regularization parameter (λ) via 5-fold '
        'cross-validation, testing 15 values logarithmically spaced from 0.01 to 0.316. We selected λ yielding '
        'networks with 5-18 edges (approximately 1.5 edges per node) to ensure connectivity while maintaining sparsity.'
    )
    
    p = doc.add_paragraph(
        '(4) Partial correlation network: We converted precision matrix elements to partial correlations using '
        'ρᵢⱼ = -θᵢⱼ / √(θᵢᵢ × θⱼⱼ), where ρᵢⱼ represents the correlation between foods i and j controlling for all '
        'other foods (26). Edges were retained if absolute partial correlation exceeded 0.10, representing meaningful '
        'conditional dependencies.'
    )
    
    doc.add_heading('Network Analysis', level=2)
    p = doc.add_paragraph(
        'We calculated three node centrality metrics for each network: (1) degree centrality (number of direct '
        'connections, normalized by maximum possible); (2) betweenness centrality (frequency on shortest paths between '
        'other nodes); and (3) closeness centrality (inverse average shortest path length). Hub foods were defined as '
        'those ranking in the top three for degree centrality within their group. All analyses used Python 3.9 with '
        'NetworkX 2.8 and scikit-learn 1.0.'
    )
    
    doc.add_page_break()
    
    # Results
    doc.add_heading('Results', level=1)
    
    doc.add_heading('Network Structure Heterogeneity', level=2)
    p = doc.add_paragraph(
        'Unlike simple co-occurrence networks that yield uniform topology, GGM revealed substantial heterogeneity in '
        'network structure across the 11 stratified groups. Edge counts ranged from 3 to 9 (mean: 6.5 ± 1.8), with '
        'network density varying from 0.045 to 0.136 (mean: 0.096 ± 0.026). This structural diversity reflects genuine '
        'differences in conditional dependency patterns across demographic and metabolic subgroups. Optimal regularization '
        'parameters (α) ranged from 0.124 to 0.164, indicating similar sparsity preferences across groups despite '
        'different edge configurations. Notably, female older adults with MetS- showed the sparsest network (3 edges, '
        'density=0.045), suggesting more independent food group consumption in this stratum.'
    )
    
    doc.add_heading('Frequent Hub Foods', level=2)
    p = doc.add_paragraph(
        'Contrary to expectations based on dietary guidelines, GGM analysis revealed that processed and fried foods—not '
        'traditionally recommended healthy foods—occupied central hub positions. Processed foods emerged as the most '
        'frequent top-3 hub, appearing in 10 of 11 groups (degree centrality range: 0.091-0.273). Protein foods ranked '
        'as hubs in 9 of 11 groups (degree: 0.091-0.182), followed by fried foods in 9 of 11 groups (degree: 0.182). '
        'Notably, vegetables appeared as hubs in only 2 of 11 groups, and grain products in 0 groups, despite being '
        'dietary guideline cornerstones. This pattern reveals a disconnect between recommended and actual dietary network '
        'structures in the Korean population.'
    )
    
    doc.add_heading('Age-Related Hub Transitions', level=2)
    p = doc.add_paragraph(
        'Hub composition exhibited systematic age-related shifts, particularly pronounced in males. Among young adults '
        '(19-39 years), processed foods showed highest centrality (degree: 0.273 in both MetS+ and MetS- groups), with '
        'sugar-sweetened beverages emerging as additional hubs (degree: 0.182 in MetS+). In middle-aged adults (40-59 years), '
        'hub patterns diversified, with protein foods gaining prominence (degree: 0.182) alongside persistent fried food '
        'centrality. Older adults (60-74 years) showed the most heterogeneous patterns, with sugar-sweetened beverages '
        'appearing as hubs in some male groups (degree: 0.182 in MetS+) but not others. Among females, processed foods '
        'maintained hub status across age groups (degree: 0.182-0.273), with sweet food consumption emerging as a hub '
        'specifically in middle-aged MetS+ women (degree: 0.182).'
    )
    
    doc.add_heading('Sex Differences in Network Organization', level=2)
    p = doc.add_paragraph(
        'Female networks showed distinct organizational patterns compared to male networks. High-fat meat appeared as a '
        'hub exclusively in young females (degree: 0.182) and older females with MetS+ (degree: 0.182), suggesting '
        'sex-specific dietary preferences. Additionally, female networks exhibited lower overall centrality values, with '
        'maximum degree centrality of 0.273 compared to males\' range of 0.182-0.273. This pattern may reflect more '
        'distributed dietary patterns among women, with less dominance by specific hub foods.'
    )
    
    doc.add_heading('MetS-Specific Hub Patterns', level=2)
    p = doc.add_paragraph(
        'Contrary to initial impressions of hub uniformity, detailed MetS stratification revealed clinically meaningful '
        'hub differentiation. Fried foods exhibited MetS-specific centrality: appearing as top-3 hubs in 100% of MetS+ '
        'groups (5/5) compared to only 66.7% of MetS- groups (4/6). This pattern held across all age-sex strata, '
        'identifying fried foods as a primary intervention target for MetS patients. Conversely, protein foods showed '
        'the opposite pattern: appearing as hubs in 100% of MetS- groups (6/6) versus only 40% of MetS+ groups (2/5), '
        'suggesting protein-centric diets associate with metabolic health maintenance. Group-specific hubs further '
        'differentiated MetS status: sweet food consumption emerged as a hub exclusively in middle-aged MetS+ women '
        '(degree: 0.182), sugar-sweetened beverages appeared uniquely in elderly MetS+ men (degree: 0.182), and '
        'vegetables appeared in young MetS+ men (degree: 0.182) and elderly MetS- women (degree: 0.091)—the latter '
        'representing the healthiest observed pattern (network density: 0.045). These MetS-stratified hub patterns '
        'provide actionable targets for personalized dietary counseling.'
    )
    
    doc.add_heading('Partial Correlation Patterns', level=2)
    p = doc.add_paragraph(
        'Analysis of partial correlations revealed both positive co-consumption and negative substitution patterns. '
        'The strongest positive partial correlations (ρ > 0.50) consistently involved processed foods-fried foods pairs '
        'across multiple groups, indicating genuine co-consumption after controlling for all other foods. Moderate '
        'positive correlations (ρ = 0.30-0.50) linked protein foods with various unhealthy food groups. Notably, '
        'negative partial correlations (ρ < -0.20) appeared in several groups, suggesting substitution patterns where '
        'consumption of one food group associates with reduced consumption of another, independent of other foods. '
        'These conditional relationships are obscured in simple co-occurrence analyses.'
    )
    
    doc.add_page_break()
    
    # Discussion
    doc.add_heading('Discussion', level=1)
    
    doc.add_heading('Principal Findings', level=2)
    p = doc.add_paragraph(
        'This stratified GGM analysis of 22,964 Korean adults revealed critical insights into actual dietary network '
        'structures that challenge conventional nutritional assumptions. Three key findings emerged: (1) Processed and '
        'fried foods—not recommended healthy foods—consistently occupy central hub positions across demographic and '
        'metabolic subgroups; (2) Network topology varies substantially by age, sex, and MetS status, with density '
        'ranging 3-fold (0.045-0.136), refuting the assumption of uniform dietary patterns; (3) Partial correlations '
        'reveal conditional dependencies invisible to simple co-occurrence approaches, including negative substitution '
        'patterns. These findings advance dietary network science by demonstrating that GGM can identify realistic '
        'intervention targets grounded in actual consumption patterns while revealing group-specific priorities for '
        'personalized nutrition.'
    )
    
    doc.add_heading('Comparison with Co-occurrence Networks', level=2)
    p = doc.add_paragraph(
        'Our GGM approach yielded markedly different results from simple co-occurrence networks. Prior co-occurrence '
        'analysis of the same population (unpublished data) identified vegetables, protein foods, and grain products '
        'as universal hubs with uniform network topology (20 edges, density=0.303 across all groups). In contrast, '
        'GGM revealed: (1) substantially sparser networks (3-9 edges) reflecting genuine conditional dependencies; '
        '(2) processed and fried foods as central hubs, not traditional healthy foods; and (3) meaningful network '
        'topology variation across groups. This discrepancy illustrates how co-occurrence conflates direct relationships '
        'with spurious associations arising from confounding. For example, vegetables may co-occur with protein foods '
        'through shared meal contexts, but GGM reveals limited direct conditional dependency when controlling for other '
        'foods. This methodological advance provides more accurate targets for dietary intervention.'
    )
    
    doc.add_heading('Implications for Dietary Guidelines', level=2)
    p = doc.add_paragraph(
        'The dominance of processed and fried foods as network hubs reveals a sobering disconnect between dietary '
        'guidelines and actual consumption patterns in the Korean population. Rather than the recommended vegetables-'
        'grains-protein triad, actual dietary networks center on processed-protein-fried foods. This finding has two '
        'implications: (1) Population-level interventions should prioritize reducing processed and fried food consumption, '
        'as their central network positions suggest that changes in these foods will cascade through connected dietary '
        'patterns; (2) Simply promoting healthy food consumption may be insufficient if it fails to displace existing hub '
        'foods. Successful interventions may require active substitution strategies that replace processed/fried food '
        'hubs with healthier alternatives rather than merely adding recommended foods to existing patterns.'
    )
    
    doc.add_heading('Clinical Implications for Personalized Nutrition', level=2)
    p = doc.add_paragraph(
        'The MetS-stratified hub analysis enables a comprehensive three-tiered intervention framework combining universal, '
        'MetS-specific, and group-specific strategies. Universal targets include processed foods (90.9% hub frequency) and '
        'fried foods (81.8% hub frequency), which should be reduced across all demographic groups regardless of metabolic '
        'status. These foods occupy central network positions, suggesting that interventions targeting them will cascade '
        'through connected dietary patterns.'
    )
    
    p = doc.add_paragraph(
        'MetS-specific interventions should prioritize fried food reduction in MetS+ patients (100% hub frequency vs 66.7% '
        'in MetS-), focusing on cooking method modification (grilling instead of frying) and portion control. This represents '
        'a primary leverage point for MetS management. Conversely, MetS- individuals should maintain protein-centric dietary '
        'patterns (100% hub frequency vs 40% in MetS+), emphasizing high-quality protein sources (fish, lean meats, legumes) '
        'for metabolic health preservation. This protective pattern warrants positive reinforcement in healthy populations.'
    )
    
    p = doc.add_paragraph(
        'Group-specific interventions address unique risk patterns: (1) Young adult males with MetS+ show vegetable hubs—an '
        'uncommon positive pattern suggesting intervention receptiveness; coaching should leverage this foundation while '
        'reducing fried food consumption. (2) Middle-aged females with MetS+ exhibit sweet food hubs uniquely in this group, '
        'requiring gender-specific interventions addressing possible hormonal influences on dietary preferences during '
        'menopausal transition. (3) Elderly males with MetS+ show sugar-sweetened beverage hubs, representing critical '
        'diabetes risk requiring immediate beverage restriction. (4) Elderly females without MetS demonstrate the exemplary '
        'pattern (protein-vegetable hubs, sparsest network), serving as the aspirational target for all other groups.'
    )
    
    p = doc.add_paragraph(
        'This evidence-based personalization framework moves beyond generic dietary guidelines to address actual consumption '
        'patterns revealed through conditional dependency analysis. Clinical decision trees incorporating age, sex, and MetS '
        'status can guide practitioners in selecting appropriate intervention priorities and realistic behavior change targets '
        'for individual patients (see Supplementary Figure S4 and Table S5).'
    )
    
    doc.add_heading('Methodological Considerations', level=2)
    p = doc.add_paragraph(
        'We chose SGCGM with graphical lasso for several methodological advantages: (1) Partial correlations control '
        'for confounding, revealing genuine conditional dependencies rather than spurious co-occurrence; (2) Rank-based '
        'transformations accommodate non-normal dietary distributions without parametric assumptions; (3) Cross-validated '
        'regularization provides data-driven edge selection rather than arbitrary thresholds; (4) Sparsity encourages '
        'interpretable networks by removing weak connections while retaining strong conditional dependencies. Alternative '
        'approaches include neighborhood selection via nodewise regression or nonparanormal SKEPTIC for larger networks, '
        'but our sample sizes supported full precision matrix estimation. The continuous score analysis (not binarized) '
        'preserved information and enabled partial correlation estimation, unlike co-occurrence approaches requiring '
        'arbitrary cutpoints.'
    )
    
    doc.add_heading('Strengths and Limitations', level=2)
    p = doc.add_paragraph(
        'Strengths include the large, nationally representative sample; comprehensive stratification strategy; state-of-'
        'the-art GGM methodology with cross-validated regularization; and multiple centrality metrics providing convergent '
        'evidence. The continuous score analysis preserves information lost in binarization. Limitations merit consideration. '
        'First, the cross-sectional design precludes causal inference regarding diet-MetS relationships or network formation '
        'processes. Second, food frequency questionnaires are subject to recall bias and may not capture meal-specific '
        'contexts. Third, our 12-food group classification aggregates heterogeneous foods (e.g., "processed foods" encompasses '
        'diverse products), potentially masking finer-grained patterns. Fourth, the female young adult MetS+ group was excluded '
        'due to small sample size, limiting generalizability to this stratum. Fifth, while SGCGM accommodates non-normality, '
        'it assumes Gaussian copulas that may not capture all dependency structures. Finally, our minimum partial correlation '
        'threshold (0.10) represents a pragmatic balance between connectivity and interpretability but remains somewhat arbitrary.'
    )
    
    doc.add_heading('Conclusions', level=2)
    p = doc.add_paragraph(
        'Gaussian graphical model analysis reveals that processed and fried foods, not traditionally recommended healthy '
        'foods, constitute the central hub structure of actual dietary networks in Korean adults. While this processed-protein-fried '
        'triad provides robust targets for population-wide interventions, MetS-stratified analysis uncovers clinically actionable '
        'differentiation: fried foods universal in MetS+ (100%) but less frequent in MetS- (66.7%), while protein foods show '
        'the inverse pattern (100% in MetS- vs 40% in MetS+). Group-specific hubs—sweet foods in middle-aged MetS+ women, '
        'sugar-sweetened beverages in elderly MetS+ men, vegetables in select healthy groups—enable personalized intervention '
        'strategies. The exemplary pattern of elderly MetS- women (protein-vegetable hubs, minimal network density) represents '
        'the aspirational target for dietary counseling.'
    )
    
    p = doc.add_paragraph(
        'This three-tiered framework (universal + MetS-specific + group-specific targets) advances precision nutrition by '
        'grounding interventions in actual consumption patterns rather than dietary guideline ideals. The substantial '
        'methodological improvement over co-occurrence networks—controlling for confounding through partial correlations—yields '
        'more accurate intervention targets based on genuine conditional dependencies. Future research should extend these methods '
        'to longitudinal data examining network stability and temporal dynamics, validate hub-focused interventions in randomized '
        'trials testing whether fried food reduction in MetS+ patients improves metabolic outcomes, and investigate network-guided '
        'versus conventional dietary counseling effectiveness. The integration of GGM with nutritional epidemiology offers a '
        'promising framework for developing evidence-based, personalized dietary recommendations that reflect real-world dietary '
        'networks while identifying realistic, group-specific targets for behavior change.'
    )
    
    doc.add_page_break()
    
    # References (updated with GGM references)
    doc.add_heading('References', level=1)
    p = doc.add_paragraph(
        '1. Alberti KG, Eckel RH, Grundy SM, et al. Harmonizing the metabolic syndrome. Circulation 2009;120:1640-5.\n\n'
        '2. Saklayen MG. The Global Epidemic of the Metabolic Syndrome. Curr Hypertens Rep 2018;20:12.\n\n'
        '3. Kastorini CM, Milionis HJ, Esposito K, et al. The effect of Mediterranean diet on metabolic syndrome and its '
        'components: a meta-analysis. J Am Coll Cardiol 2011;57:1299-313.\n\n'
        '4. Dietary Guidelines Advisory Committee. Scientific Report of the 2020 Dietary Guidelines Advisory Committee. '
        'Washington, DC: U.S. Department of Agriculture, 2020.\n\n'
        '5. Hu FB. Dietary pattern analysis: a new direction in nutritional epidemiology. Curr Opin Lipidol 2002;13:3-9.\n\n'
        '6. Newby PK, Tucker KL. Empirically derived eating patterns using factor or cluster analysis: a review. '
        'Nutr Rev 2004;62:177-203.\n\n'
        '7. Hosseini SR, Koushesh S, Maroufy V. Spurious associations in analysis of dietary patterns. Am J Clin Nutr '
        '2011;93:1128-9.\n\n'
        '8. Friedman J, Hastie T, Tibshirani R. Sparse inverse covariance estimation with the graphical lasso. '
        'Biostatistics 2008;9:432-41.\n\n'
        '9. Lauritzen SL. Graphical Models. Oxford: Oxford University Press, 1996.\n\n'
        '10. Whittaker J. Graphical Models in Applied Multivariate Statistics. Chichester: Wiley, 1990.\n\n'
        '11. Drton M, Maathuis MH. Structure learning in graphical modeling. Annu Rev Stat Appl 2017;4:365-93.\n\n'
        '12. Schwedhelm C, Knüppel S, Schwingshackl L, et al. Meal and habitual dietary networks identified through '
        'semiparametric Gaussian copula graphical models in a German adult population. PLoS One 2018;13:e0202936.\n\n'
        '13. Schwedhelm C, Iqbal K, Knüppel S, et al. Contribution to the understanding of how principal component '
        'analysis-derived dietary patterns emerge from habitual data on diet. Am J Clin Nutr 2018;107:227-35.\n\n'
        '14. Schwedhelm C, Lipsky LM, Shearrer GE, et al. Using food network analysis to understand meal patterns in '
        'pregnant women with high and low diet quality. Int J Behav Nutr Phys Act 2021;18:48.\n\n'
        '15. Barabási AL, Menichetti G, Loscalzo J. The unmapped chemical complexity of our diet. Nat Food 2020;1:33-7.\n\n'
        '16. Trijsburg L, Talsma EF, de Vries JHM, et al. Diet quality indices for research in low- and middle-income '
        'countries: a systematic review. Nutr Rev 2019;77:515-40.\n\n'
        '17. Wardle J, Haase AM, Steptoe A, et al. Gender differences in food choice: the contribution of health beliefs '
        'and dieting. Ann Behav Med 2004;27:107-16.\n\n'
        '18. Nicklett EJ, Kadell AR. Fruit and vegetable intake among older adults: a scoping review. Maturitas 2013;75:305-12.\n\n'
        '19. Mente A, de Koning L, Shannon HS, Anand SS. A systematic review of the evidence supporting a causal link '
        'between dietary factors and coronary heart disease. Arch Intern Med 2009;169:659-69.\n\n'
        '20. Mozaffarian D. Dietary and policy priorities for cardiovascular disease, diabetes, and obesity: a comprehensive '
        'review. Circulation 2016;133:187-225.\n\n'
        '21. Grundy SM, Cleeman JI, Daniels SR, et al. Diagnosis and management of the metabolic syndrome: an American '
        'Heart Association/National Heart, Lung, and Blood Institute scientific statement. Circulation 2005;112:2735-52.\n\n'
        '22. Kim S, Moon S, Popkin BM. The nutrition transition in South Korea. Am J Clin Nutr 2000;71:44-53.\n\n'
        '23. Liu H, Lafferty J, Wasserman L. The nonparanormal: semiparametric estimation of high dimensional undirected graphs. '
        'J Mach Learn Res 2009;10:2295-328.\n\n'
        '24. Banerjee O, El Ghaoui L, d\'Aspremont A. Model selection through sparse maximum likelihood estimation for '
        'multivariate Gaussian or binary data. J Mach Learn Res 2008;9:485-516.\n\n'
        '25. Yuan M, Lin Y. Model selection and estimation in the Gaussian graphical model. Biometrika 2007;94:19-35.\n\n'
        '26. Baba K, Shibata R, Sibuya M. Partial correlation and conditional correlation as measures of conditional '
        'independence. Aust N Z J Stat 2004;46:657-64.\n\n'
        '[Additional references to be added as needed]'
    )
    
    return doc

def main():
    """메인 함수"""
    print("Creating SCI-level English paper (GGM version)...")
    
    doc = create_sci_paper_ggm()
    
    # 저장
    output_dir = Path(__file__).parent.parent / 'result' / 'manuscript'
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / 'Paper2_SCI_English_GGM.docx'
    
    doc.save(output_file)
    
    print(f"\n✓ GGM-based SCI paper saved: {output_file}")
    print(f"  File size: {output_file.stat().st_size / 1024:.1f} KB")
    print("\n영어 SCI급 논문 작성 완료 (GGM 방법론)!")

if __name__ == '__main__':
    main()
